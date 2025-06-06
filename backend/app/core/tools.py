import requests
import html
import re
import json
import trafilatura
from typing import Optional, Any, Dict, List
import logging
from langchain_openai import ChatOpenAI
import os
from dotenv import load_dotenv, find_dotenv
from groq import Groq
from docx import Document
from openai import OpenAI
load_dotenv(find_dotenv())

def search_serper(query: str, api_key: str, num_results: int = 3) -> List[str]:
    """Search using Serper API and return top URLs."""
    url = "https://google.serper.dev/search"
    headers = {
        "X-API-KEY": api_key,
        "Content-Type": "application/json"
    }
    payload = {"q": query}
    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        results = response.json().get("organic", [])[:num_results]
        print(f"[INFO]: Results from serper: ")
        return [r["link"] for r in results if "link" in r]
    except Exception as e:
        logging.error(f"Serper search failed: {e}")
        return []

def clean_scraped_text(raw_text: str) -> str:
    """Clean and normalize scraped text."""
    if not raw_text:
        return ""
    text = html.unescape(raw_text)
    text = text.encode().decode('unicode_escape')
    text = text.replace("\\'", "'")
    text = re.sub(r'\|[-\s]*\|', '', text)
    text = re.sub(r'\|.*?\|', '', text)
    text = re.sub(r'\n\s*', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = text.encode("ascii", errors="ignore").decode()
    print(f"[INFO]: Text Cleaned Successfully!")
    return text.strip()

def scrape_and_clean(url: str) -> str:
    """Scrape and clean text from a URL."""
    try:
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            clean_text = trafilatura.extract(downloaded)
            clean_text = clean_scraped_text(clean_text)
            print(f"[INFO]: Scraped text cleaned!!")
            return clean_text or ""
        
        return ""
    except Exception as e:
        logging.error(f"Scraping failed for {url}: {e}")
        return ""

def summarize_text(text: str, user_query: str, model="gpt-4.1-nano") -> str:
    prompt = f"""Summarize the following content in a concise, informative way:\n\n{text}. Only keep the information required by the user. User query: {user_query}.
    <Instruction> 1. keep only the necessary information in the summary. Keep the summary within 150 words. """ 
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    print(f"[INFO]: Summarized the text! summary", response)
    return response.choices[0].message.content.strip()

def map_to_schema(summary: str, user_query: str, model="gpt-4.1-nano") -> dict:
    schema_prompt = f"""
You are a business research assistant. Analyze and Extract structured info from the text below based on the user query.
User Query:
{user_query}

Text:
{summary}

Return a JSON with possible fields. 
<Instruction>
1. Respond only with JSON.
2. Keep the fields required by the user.
"""
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": schema_prompt}],
        temperature=0.2
    )
    print("response from map to schema", response)
    try:
        return json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        return {"raw": response.choices[0].message.content}
 

def extract_json_block(response: dict) -> dict:
    raw_text = response.get("raw", "")

    clean_json_str = re.sub(r'^```[\w]*\n?', '', raw_text.strip())
    clean_json_str = re.sub(r'```$', '', clean_json_str.strip())

    try:
        return json.loads(clean_json_str)
    except json.JSONDecodeError as e:
        print("Failed to decode JSON:", e)
        return {}
    
def summarize_or_filter_json(
    json_data: Dict[str, Any],
    user_query: str,
    summarize: bool = True,
    model="gpt-4.1-nano" # Pass your LLM client if available
) -> Dict[str, Any]:
    """
    Summarizes or filters a JSON response to keep only essential info based on the user query.
    
    Args:
        json_data: The original JSON data (e.g., from map_to_schema_node).
        user_query: The user's original query (for context).
        summarize: If True, summarize the data; if False, filter by relevance.
        llm_client: Optional LLM client for advanced summarization/filtering.
    
    Returns:
        Filtered or summarized JSON.
    """
    if not json_data or "results" not in json_data:
        return json_data


    prompt = f"""
        You are a business research assistant. The user asked: "{user_query}".
        Here is the data to process:
        {json.dumps(json_data["results"], indent=2)}
        
        Please return a JSON array with only the most relevant or summarized info,
        based on the user query. Remove unnecessary or empty keys from the final json return a clean response combining all results. 
        """
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2
    )
    try:
        return json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        return {"raw": response.choices[0].message.content}

def json_to_docx(json_data: dict, user_query: str, filename: str):
    """
    Convert JSON to a Word document.
    """
    doc = Document()
    doc.add_heading(f"Report for query: {user_query}", level=1)
    if not json_data or "results" not in json_data or not json_data["results"]:
        doc.add_paragraph("No results found.")
        doc.save(filename)
        return

    for i, result in enumerate(json_data["results"], 1):
        doc.add_heading(f"Result {i}", level=2)
        for key, value in result.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.add_paragraph("")

    doc.save(filename)
